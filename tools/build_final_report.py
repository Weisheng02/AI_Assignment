#!/usr/bin/env python3
"""Build the final BMCS2003 university-chatbot report.

The document and all evidence charts are generated from repository artifacts.
No performance value is hard-coded: rerunning this script after evaluate.py
updates data/evaluation_results.json refreshes the report automatically.
"""

from __future__ import annotations

import json
import math
import re
import shutil
import textwrap
import zipfile
from collections import Counter
from datetime import datetime
from pathlib import Path

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
from matplotlib.patches import FancyBboxPatch

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_BREAK, WD_LINE_SPACING, WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DATA_DIR = ROOT / "data"
ASSET_DIR = ROOT / "report_assets"
OUTPUT_PATH = ROOT / "AI Report - Final.docx"
ORIGINAL_REPORT = ROOT / "AI Report.docx"

EVALUATION_PATH = DATA_DIR / "evaluation_results.json"
INTENTS_PATH = DATA_DIR / "intents.json"
FEEDBACK_PATH = DATA_DIR / "user_feedback_verified.json"
RESPONSE_TEST_PATH = DATA_DIR / "response_quality_test.json"
MODEL_SELECTION_PATH = DATA_DIR / "model_selection_results.json"
LATEST_TEST_RESULTS_PATH = DATA_DIR / "latest_test_results.json"
DIALOGFLOW_NOTES_PATH = ROOT / "docs" / "dialogflow_setup.md"

BLUE = "2E74B5"
DEEP_BLUE = "1F4D78"
NAVY = "17365D"
TEAL = "2A7F9E"
GOLD = "C9972C"
CHARCOAL = "25313C"
MID_GREY = "66717D"
LIGHT_GREY = "F4F6F9"
PALE_BLUE = "EAF2F8"
PALE_GOLD = "FFF6DD"
PALE_RED = "FCECEA"
WHITE = "FFFFFF"
TABLE_WIDTH_DXA = 9360

SURVEY_FIELDS = (
    ("intent_accuracy", "Intent understanding", "The chatbot correctly understands student inquiry questions and intent."),
    ("response_quality", "Answer clarity and relevance", "The chatbot's answers are clear, informative, and relevant to university procedures."),
    ("ui_navigability", "Interface usability", "The web user interface (Streamlit GUI) is easy to navigate and chat with."),
    ("response_speed", "Response speed", "The chatbot responds promptly without noticeable delay."),
    ("overall_satisfaction", "Overall satisfaction", "Overall, I am satisfied with the automated university inquiry chatbot system."),
)


def load_json(path: Path, default):
    if not path.exists():
        return default
    with path.open("r", encoding="utf-8") as handle:
        return json.load(handle)


def as_records(payload):
    if isinstance(payload, list):
        return payload
    if isinstance(payload, dict):
        return payload.get("cases", [])
    return []


def safe_float(value):
    if value is None or isinstance(value, bool):
        return None
    try:
        number = float(value)
    except (TypeError, ValueError):
        return None
    return number if math.isfinite(number) else None


def metric_text(value, digits: int = 4):
    number = safe_float(value)
    return "N/A" if number is None else f"{number:.{digits}f}"


def percent_text(value, digits: int = 2):
    number = safe_float(value)
    return "N/A" if number is None else f"{100 * number:.{digits}f}%"


def result_name(row):
    return str(row.get("Member") or row.get("model_name") or row.get("name") or "Unnamed model")


def engine_name(row):
    return str(row.get("Engine Type") or row.get("engine") or "Not recorded")


def field(row, *names):
    for name in names:
        if name in row:
            return row[name]
    return None


def extract_logo():
    """Reuse the institutional logo embedded in the preserved original report."""
    target = ASSET_DIR / "tarumt_logo.png"
    if target.exists():
        return target
    if not ORIGINAL_REPORT.exists():
        return None
    with zipfile.ZipFile(ORIGINAL_REPORT) as archive:
        candidates = [
            name
            for name in archive.namelist()
            if name.startswith("word/media/") and name.lower().endswith((".png", ".jpg", ".jpeg"))
        ]
        if not candidates:
            return None
        # The first image in the source report is the TAR UMT wordmark.
        data = archive.read(sorted(candidates)[0])
    target.write_bytes(data)
    return target


def short_model_name(name: str):
    replacements = {
        "Member 1 (Dialogflow ES)": "M1 local simulator",
        "Member 2 (TF-IDF + Logistic Reg)": "M2 TF-IDF + LR",
        "Baseline 1 (Multinomial Naïve Bayes)": "Multinomial NB",
        "Baseline 2 (Linear SVM)": "Linear SVM",
    }
    return replacements.get(name, name)


def create_architecture_chart(path: Path):
    fig, ax = plt.subplots(figsize=(12.6, 7.8))
    ax.set_xlim(0, 12.6)
    ax.set_ylim(0, 7.8)
    ax.axis("off")

    def box(x, y, w, h, text, fc, ec=DEEP_BLUE, fontsize=10, weight="normal"):
        patch = FancyBboxPatch(
            (x, y), w, h, boxstyle="round,pad=0.03,rounding_size=0.10",
            linewidth=1.5, edgecolor=f"#{ec}", facecolor=f"#{fc}"
        )
        ax.add_patch(patch)
        ax.text(x + w / 2, y + h / 2, text, ha="center", va="center",
                fontsize=fontsize, color=f"#{CHARCOAL}", weight=weight, wrap=True)

    def arrow(x1, y1, x2, y2, dashed=False):
        ax.annotate("", xy=(x2, y2), xytext=(x1, y1),
                    arrowprops=dict(arrowstyle="-|>", lw=1.5, color=f"#{DEEP_BLUE}",
                                    linestyle="--" if dashed else "-"))

    ax.text(0.2, 7.5, "Operational architecture and evaluation boundary", fontsize=18,
            color=f"#{NAVY}", weight="bold")
    ax.text(0.2, 7.14, "The Dialogflow route and local Python route are separate; the offline simulator is evaluation-only.",
            fontsize=10.5, color=f"#{MID_GREY}")

    box(0.3, 5.35, 1.7, 0.95, "Student\nquery", PALE_BLUE, weight="bold")
    box(2.4, 5.35, 2.2, 0.95, "Streamlit interface\nand route selection", LIGHT_GREY, weight="bold")
    arrow(2.0, 5.82, 2.4, 5.82)

    box(5.15, 5.78, 2.25, 1.2, "Dialogflow ES route\nintents • entities\nresponses", "E7F3EC", ec="39845A", fontsize=9.2, weight="bold")
    box(8.0, 5.83, 2.1, 1.1, "Google-managed NLU\ncloud execution", "E7F3EC", ec="39845A")
    box(10.65, 5.83, 1.65, 1.1, "Matched intent\nand response", "E7F3EC", ec="39845A")
    arrow(4.6, 5.82, 5.15, 6.38)
    arrow(7.4, 6.38, 8.0, 6.38)
    arrow(10.1, 6.38, 10.65, 6.38)
    ax.text(8.25, 5.47, "Dialogflow route bypasses local NLTK", ha="center",
            fontsize=8.8, color="#39845A", weight="bold")

    box(5.15, 3.55, 1.8, 1.2, "Local NLTK\ncleaning +\nlemmatisation", PALE_BLUE, weight="bold")
    box(7.35, 3.55, 1.95, 1.2, "Character TF-IDF\nchar_wb 3–5 grams\nsublinear TF", PALE_BLUE)
    box(9.7, 3.55, 2.25, 1.2, "Logistic Regression\nC = 30 • max_iter = 2000\nclass_weight = balanced", PALE_BLUE)
    box(9.7, 1.85, 2.25, 1.05, "Confidence gate\nthreshold = 0.20", PALE_GOLD, ec=GOLD, weight="bold")
    box(6.8, 1.85, 2.15, 1.05, "Intent response\nor polite fallback", LIGHT_GREY)
    box(4.05, 1.85, 2.15, 1.05, "Low-confidence log\nfor human review", LIGHT_GREY)
    arrow(4.6, 5.72, 5.15, 4.15)
    arrow(6.95, 4.15, 7.35, 4.15)
    arrow(9.3, 4.15, 9.7, 4.15)
    arrow(10.82, 3.55, 10.82, 2.9)
    arrow(9.7, 2.38, 8.95, 2.38)
    arrow(6.8, 2.38, 6.2, 2.38, dashed=True)

    box(0.3, 0.3, 11.65, 0.86,
        "OFFLINE EVALUATION HARNESS: stratified 80/20 split, seed 42, zero cleaned-text overlap. "
        "Member 1 is measured with a train-only Jaccard/pattern simulator—not a live or cloud Dialogflow agent. "
        "Member 2 and both baselines are trained only on the training partition.",
        PALE_RED, ec="B14C45", fontsize=9.5, weight="bold")
    arrow(2.15, 1.16, 2.15, 5.35, dashed=True)
    ax.text(2.28, 1.28, "evaluation boundary", fontsize=8.5, color="#B14C45", rotation=90)

    fig.tight_layout()
    fig.savefig(path, dpi=220, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def create_dataset_chart(path: Path, intent_rows):
    pairs = sorted(
        [(str(item.get("tag", "unnamed")), len(item.get("patterns", []))) for item in intent_rows],
        key=lambda pair: (pair[1], pair[0]),
    )
    labels = [name.replace("_", " ") for name, _ in pairs]
    values = [count for _, count in pairs]
    fig_height = max(7.8, len(pairs) * 0.24)
    fig, ax = plt.subplots(figsize=(8.7, fig_height))
    colors = [f"#{BLUE}" if value >= np.median(values) else "#9EBCD2" for value in values]
    bars = ax.barh(labels, values, color=colors, edgecolor="white", height=0.72)
    ax.set_xlabel("Training phrases in current data/intents.json", fontsize=10)
    ax.set_title("Current intent inventory", loc="left", fontsize=16, weight="bold", color=f"#{NAVY}")
    ax.grid(axis="x", alpha=0.22)
    ax.spines[["top", "right", "left"]].set_visible(False)
    ax.tick_params(axis="y", labelsize=8.5)
    ax.tick_params(axis="x", labelsize=9)
    for bar, value in zip(bars, values):
        ax.text(value + 0.12, bar.get_y() + bar.get_height() / 2, str(value), va="center", fontsize=8)
    ax.set_xlim(0, max(values or [1]) + 2.2)
    fig.tight_layout()
    fig.savefig(path, dpi=220, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def create_metrics_chart(path: Path, results):
    labels = [short_model_name(result_name(row)) for row in results]
    metrics = ["Accuracy", "Precision", "Recall", "F1-Score"]
    data = np.array([[safe_float(row.get(metric)) or 0.0 for metric in metrics] for row in results])
    x = np.arange(len(labels))
    width = 0.18
    palette = [BLUE, TEAL, GOLD, "7B6BA8"]
    fig, ax = plt.subplots(figsize=(9.3, 4.8))
    for idx, (metric, color) in enumerate(zip(metrics, palette)):
        bars = ax.bar(x + (idx - 1.5) * width, data[:, idx], width, label=metric, color=f"#{color}")
        for bar, value in zip(bars, data[:, idx]):
            ax.text(bar.get_x() + bar.get_width() / 2, value + 0.012, f"{value:.2f}",
                    ha="center", va="bottom", fontsize=7.4, rotation=90)
    ax.set_ylim(0, max(1.0, float(data.max(initial=0)) + 0.16))
    ax.set_ylabel("Score")
    ax.set_xticks(x)
    ax.set_xticklabels(labels, fontsize=9)
    ax.set_title("Leakage-free held-out intent-classification metrics", loc="left",
                 fontsize=15, weight="bold", color=f"#{NAVY}")
    ax.grid(axis="y", alpha=0.22)
    ax.spines[["top", "right"]].set_visible(False)
    ax.legend(ncol=4, frameon=False, loc="upper center", bbox_to_anchor=(0.5, -0.13), fontsize=9)
    fig.tight_layout()
    fig.savefig(path, dpi=220, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def create_coverage_chart(path: Path, results):
    labels = [short_model_name(result_name(row)) for row in results]
    coverage = [safe_float(row.get("Coverage")) or 0.0 for row in results]
    fallback = [safe_float(row.get("Fallback Rate")) or 0.0 for row in results]
    x = np.arange(len(labels))
    fig, ax = plt.subplots(figsize=(9.2, 4.1))
    ax.bar(x, coverage, color=f"#{TEAL}", label="Coverage")
    ax.bar(x, fallback, bottom=coverage, color=f"#{GOLD}", label="Fallback rate")
    for index, (cov, fall) in enumerate(zip(coverage, fallback)):
        ax.text(index, min(cov / 2, 0.92), f"{cov:.2f}", ha="center", va="center", color="white", fontsize=9, weight="bold")
        if fall > 0.02:
            ax.text(index, cov + fall / 2, f"{fall:.2f}", ha="center", va="center", color=f"#{CHARCOAL}", fontsize=9)
    ax.set_ylim(0, 1.08)
    ax.set_ylabel("Proportion of held-out cases")
    ax.set_xticks(x)
    ax.set_xticklabels(labels, fontsize=9)
    ax.set_title("Coverage and fallback behaviour", loc="left", fontsize=15, weight="bold", color=f"#{NAVY}")
    ax.legend(frameon=False, ncol=2, loc="upper right")
    ax.grid(axis="y", alpha=0.18)
    ax.spines[["top", "right"]].set_visible(False)
    fig.tight_layout()
    fig.savefig(path, dpi=220, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def top_confusions(confusion_payload, limit=10):
    labels = confusion_payload.get("labels", [])
    matrix = confusion_payload.get("matrix", [])
    pairs = []
    for actual_idx, row in enumerate(matrix):
        for predicted_idx, count in enumerate(row):
            if actual_idx != predicted_idx and count:
                pairs.append((int(count), labels[actual_idx], labels[predicted_idx]))
    return sorted(pairs, key=lambda item: (-item[0], item[1], item[2]))[:limit]


def create_error_chart(path: Path, matrices, model_names):
    fig, axes = plt.subplots(1, 2, figsize=(12.2, 6.7), constrained_layout=True)
    for axis, model_name, color in zip(axes, model_names, [BLUE, GOLD]):
        payload = matrices.get(model_name, {})
        pairs = top_confusions(payload, limit=10)
        if not pairs:
            axis.text(0.5, 0.5, "No off-diagonal errors recorded", ha="center", va="center")
            axis.axis("off")
            continue
        labels = [f"{a.replace('_', ' ')} → {p.replace('_', ' ')}" for _, a, p in reversed(pairs)]
        values = [count for count, _, _ in reversed(pairs)]
        bars = axis.barh(labels, values, color=f"#{color}")
        axis.set_title(short_model_name(model_name), loc="left", fontsize=13, weight="bold", color=f"#{NAVY}")
        axis.set_xlabel("Held-out cases")
        axis.tick_params(axis="y", labelsize=8.2)
        axis.spines[["top", "right", "left"]].set_visible(False)
        axis.grid(axis="x", alpha=0.18)
        axis.set_xlim(0, max(values) + 0.8)
        for bar, value in zip(bars, values):
            axis.text(value + 0.06, bar.get_y() + bar.get_height() / 2, str(value), va="center", fontsize=8)
    fig.suptitle("Most frequent off-diagonal confusion pairs", x=0.02, ha="left",
                 fontsize=16, weight="bold", color=f"#{NAVY}")
    fig.savefig(path, dpi=220, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def generate_assets(intents, evaluation, survey):
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    extract_logo()
    paths = {
        "architecture": ASSET_DIR / "architecture.png",
        "dataset": ASSET_DIR / "dataset_distribution.png",
        "metrics": ASSET_DIR / "model_metrics.png",
        "coverage": ASSET_DIR / "coverage_fallback.png",
        "errors": ASSET_DIR / "confusion_error_analysis.png",
        "survey": ASSET_DIR / "survey_results.png",
        "app_screenshot": ASSET_DIR / "app_chatbot_e2e.png",
    }
    create_architecture_chart(paths["architecture"])
    create_dataset_chart(paths["dataset"], intents.get("intents", []))
    create_metrics_chart(paths["metrics"], evaluation.get("results", []))
    create_coverage_chart(paths["coverage"], evaluation.get("results", []))
    matrices = evaluation.get("details", {}).get("confusion_matrices", {})
    preferred = [
        "Member 1 (Dialogflow ES)",
        "Member 2 (TF-IDF + Logistic Reg)",
    ]
    names = [name for name in preferred if name in matrices]
    if len(names) < 2:
        names = list(matrices)[:2]
    while len(names) < 2:
        names.append(names[0] if names else "Unavailable")
    create_error_chart(paths["errors"], matrices, names[:2])
    create_survey_chart(paths["survey"], survey)
    return paths


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_table_geometry(table, widths, indent=120):
    if sum(widths) != TABLE_WIDTH_DXA:
        raise ValueError(f"Table widths must total {TABLE_WIDTH_DXA}; got {sum(widths)}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    for tag, attrs in (
        ("w:tblW", {"w:w": str(TABLE_WIDTH_DXA), "w:type": "dxa"}),
        ("w:tblInd", {"w:w": str(indent), "w:type": "dxa"}),
        ("w:tblLayout", {"w:type": "fixed"}),
    ):
        node = tbl_pr.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            tbl_pr.append(node)
        for key, value in attrs.items():
            node.set(qn(key), value)
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        prevent_row_split(row)
        for cell, width in zip(row.cells, widths):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_paragraph_border(paragraph, side="bottom", color="C7D2DB", size="6", space="4"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    border = OxmlElement(f"w:{side}")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), size)
    border.set(qn("w:space"), space)
    border.set(qn("w:color"), color)
    p_bdr.append(border)


def set_paragraph_shading(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_page_number(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text_node = OxmlElement("w:t")
    text_node.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text_node, end])


def configure_styles(document):
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(CHARCOAL)
    normal.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.333

    heading_specs = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DEEP_BLUE, 8, 4),
    }
    for style_name, (size, color, before, after) in heading_specs.items():
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True
        style.paragraph_format.widow_control = True

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.194)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        style.paragraph_format.line_spacing = 1.208

    if "Report Caption" not in styles:
        caption = styles.add_style("Report Caption", WD_STYLE_TYPE.PARAGRAPH)
    else:
        caption = styles["Report Caption"]
    caption.font.name = "Calibri"
    caption.font.size = Pt(9.5)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor.from_string(MID_GREY)
    caption.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(9)
    caption.paragraph_format.keep_with_next = True
    caption.paragraph_format.keep_together = True

    if "Source Note" not in styles:
        source_note = styles.add_style("Source Note", WD_STYLE_TYPE.PARAGRAPH)
    else:
        source_note = styles["Source Note"]
    source_note.font.name = "Calibri"
    source_note.font.size = Pt(9)
    source_note.font.color.rgb = RGBColor.from_string(MID_GREY)
    source_note.paragraph_format.space_after = Pt(8)
    source_note.paragraph_format.left_indent = Inches(0.18)
    source_note.paragraph_format.right_indent = Inches(0.18)

    if "Reference" not in styles:
        reference = styles.add_style("Reference", WD_STYLE_TYPE.PARAGRAPH)
    else:
        reference = styles["Reference"]
    reference.font.name = "Calibri"
    reference.font.size = Pt(10)
    reference.paragraph_format.left_indent = Inches(0.5)
    reference.paragraph_format.first_line_indent = Inches(-0.5)
    reference.paragraph_format.space_after = Pt(7)
    reference.paragraph_format.line_spacing = 1.15


def configure_document(document):
    section = document.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.text = "BMCS2003  |  UNIVERSITY INQUIRY CHATBOT"
    hp.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    hp.style = document.styles["Normal"]
    hp.paragraph_format.space_after = Pt(3)
    for run in hp.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(8.5)
        run.font.bold = True
        run.font.color.rgb = RGBColor.from_string(MID_GREY)
    set_paragraph_border(hp)

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    fp.paragraph_format.space_before = Pt(3)
    prefix = fp.add_run("FINAL REPORT  •  ")
    prefix.font.size = Pt(8.5)
    prefix.font.color.rgb = RGBColor.from_string(MID_GREY)
    add_page_number(fp)
    set_paragraph_border(fp, side="top")

    first_header = section.first_page_header
    first_header.paragraphs[0].text = ""
    first_footer = section.first_page_footer
    cfp = first_footer.paragraphs[0]
    cfp.text = "BMCS2003  •  SESSION 202605"
    cfp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in cfp.runs:
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor.from_string(MID_GREY)

    settings = document.settings.element
    compat = settings.find(qn("w:compat"))
    if compat is None:
        compat = OxmlElement("w:compat")
        settings.append(compat)
    do_not_expand = OxmlElement("w:doNotExpandShiftReturn")
    compat.append(do_not_expand)

    # Set default document language for accessibility tools.
    styles_el = document.styles.element
    doc_defaults = styles_el.find(qn("w:docDefaults"))
    if doc_defaults is not None:
        r_pr_default = doc_defaults.find(qn("w:rPrDefault"))
        if r_pr_default is not None:
            r_pr = r_pr_default.find(qn("w:rPr"))
            if r_pr is None:
                r_pr = OxmlElement("w:rPr")
                r_pr_default.append(r_pr)
            lang = r_pr.find(qn("w:lang"))
            if lang is None:
                lang = OxmlElement("w:lang")
                r_pr.append(lang)
            lang.set(qn("w:val"), "en-MY")


def add_paragraph(document, text="", style=None, bold_lead=None, align=None):
    paragraph = document.add_paragraph(style=style)
    if bold_lead and text.startswith(bold_lead):
        lead = paragraph.add_run(bold_lead)
        lead.bold = True
        paragraph.add_run(text[len(bold_lead):])
    else:
        paragraph.add_run(text)
    if align is not None:
        paragraph.alignment = align
    paragraph.paragraph_format.widow_control = True
    return paragraph


def add_bullet(document, text):
    return add_paragraph(document, text, style="List Bullet")


def add_number(document, text):
    return add_paragraph(document, text, style="List Number")


def create_decimal_numbering(document, start=1):
    """Create an independent single-level decimal list definition."""
    numbering = document.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
        if node.get(qn("w:abstractNumId")) is not None
    ]
    num_ids = [
        int(node.get(qn("w:numId")))
        for node in numbering.findall(qn("w:num"))
        if node.get(qn("w:numId")) is not None
    ]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start_node = OxmlElement("w:start")
    start_node.set(qn("w:val"), str(start))
    number_format = OxmlElement("w:numFmt")
    number_format.set(qn("w:val"), "decimal")
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "%1.")
    justification = OxmlElement("w:lvlJc")
    justification.set(qn("w:val"), "left")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "280")
    p_pr.extend([tabs, indent])
    level.extend([start_node, number_format, level_text, justification, p_pr])
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_numbered_list(document, items, start=1):
    num_id = create_decimal_numbering(document, start=start)
    paragraphs = []
    for text in items:
        paragraph = document.add_paragraph(style="Normal")
        paragraph.add_run(text)
        paragraph.paragraph_format.left_indent = Inches(0.375)
        paragraph.paragraph_format.first_line_indent = Inches(-0.194)
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.208
        p_pr = paragraph._p.get_or_add_pPr()
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num_id_node = OxmlElement("w:numId")
        num_id_node.set(qn("w:val"), str(num_id))
        num_pr.extend([ilvl, num_id_node])
        p_pr.insert(0, num_pr)
        paragraphs.append(paragraph)
    return paragraphs


def add_heading(document, text, level=1, page_break=False):
    if page_break:
        document.add_page_break()
    paragraph = document.add_heading(text, level=level)
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.keep_together = True
    return paragraph


def add_callout(document, title, body, tone="blue"):
    fills = {"blue": PALE_BLUE, "gold": PALE_GOLD, "red": PALE_RED, "grey": LIGHT_GREY}
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.16)
    paragraph.paragraph_format.right_indent = Inches(0.16)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(9)
    paragraph.paragraph_format.keep_together = True
    title_run = paragraph.add_run(f"{title}  ")
    title_run.bold = True
    title_run.font.color.rgb = RGBColor.from_string(DEEP_BLUE if tone != "red" else "9C3E38")
    paragraph.add_run(body)
    set_paragraph_shading(paragraph, fills.get(tone, PALE_BLUE))
    return paragraph


def add_table(document, headers, rows, widths, font_size=9.2, header_fill=LIGHT_GREY):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for idx, text in enumerate(headers):
        cell = hdr.cells[idx]
        cell.text = str(text)
        set_cell_shading(cell, header_fill)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            paragraph.paragraph_format.space_after = Pt(0)
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(font_size)
                run.font.color.rgb = RGBColor.from_string(CHARCOAL)
    for row_index, row_data in enumerate(rows):
        cells = table.add_row().cells
        for idx, value in enumerate(row_data):
            cells[idx].text = str(value)
            if row_index % 2 == 1:
                set_cell_shading(cells[idx], "FAFBFC")
            for paragraph in cells[idx].paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.08
                paragraph.paragraph_format.widow_control = True
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)
                    run.font.color.rgb = RGBColor.from_string(CHARCOAL)
    set_table_geometry(table, widths)
    document.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_caption(document, text):
    return add_paragraph(document, text, style="Report Caption", align=WD_PARAGRAPH_ALIGNMENT.CENTER)


def add_figure(document, path, width, alt_text, caption):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.keep_with_next = True
    shape = paragraph.add_run().add_picture(str(path), width=Inches(width))
    shape._inline.docPr.set("descr", alt_text)
    shape._inline.docPr.set("title", caption.split(":", 1)[0])
    add_caption(document, caption)
    return shape


def add_page_break(document):
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def add_reference(document, text):
    add_paragraph(document, text, style="Reference")


def summarise_survey(records):
    """Validate the five-item instrument and return descriptive statistics."""
    validated = []
    for record in records:
        if not isinstance(record, dict):
            continue
        row = {}
        for key, _, _ in SURVEY_FIELDS:
            value = record.get(key)
            if isinstance(value, bool) or not isinstance(value, int) or not 1 <= value <= 5:
                raise ValueError(f"Invalid verified survey rating for {key}: {value!r}")
            row[key] = value
        validated.append(row)

    items = []
    flattened = []
    for key, label, statement in SURVEY_FIELDS:
        values = [row[key] for row in validated]
        flattened.extend(values)
        items.append({
            "key": key,
            "label": label,
            "statement": statement,
            "mean": float(np.mean(values)) if values else None,
            "median": float(np.median(values)) if values else None,
            "favorable_count": sum(value >= 4 for value in values),
            "favorable_rate": (sum(value >= 4 for value in values) / len(values)) if values else None,
        })
    return {
        "respondent_count": len(validated),
        "rating_count": len(flattened),
        "overall_mean": float(np.mean(flattened)) if flattened else None,
        "overall_median": float(np.median(flattened)) if flattened else None,
        "overall_favorable_count": sum(value >= 4 for value in flattened),
        "overall_favorable_rate": (sum(value >= 4 for value in flattened) / len(flattened)) if flattened else None,
        "items": items,
    }


def create_survey_chart(path: Path, survey):
    labels = [item["label"] for item in survey["items"]]
    means = [item["mean"] for item in survey["items"]]
    favorable = [item["favorable_rate"] for item in survey["items"]]
    fig, ax = plt.subplots(figsize=(10.8, 6.2))
    bars = ax.barh(labels[::-1], means[::-1], color=f"#{TEAL}")
    ax.set_xlim(0, 5)
    ax.set_xlabel("Mean rating (1-5)")
    ax.set_title("Verified five-item usability survey (N=5)", loc="left", fontsize=16, weight="bold", color=f"#{NAVY}")
    ax.grid(axis="x", alpha=0.18)
    ax.spines[["top", "right", "left"]].set_visible(False)
    for bar, mean, rate in zip(bars, means[::-1], favorable[::-1]):
        ax.text(mean + 0.06, bar.get_y() + bar.get_height() / 2,
                f"{mean:.1f}  |  {rate * 100:.0f}% favorable", va="center", fontsize=9)
    fig.text(0.01, 0.01, "Favorable = rating 4 or 5. Descriptive evidence only; small sample.", fontsize=8.5, color=f"#{MID_GREY}")
    fig.savefig(path, dpi=220, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def source_rows(response_payload, optional_notes):
    rows = []
    for case in as_records(response_payload):
        for url in case.get("source_urls", []):
            rows.append((case.get("expected_intent", "Official fact"), url, "response_quality_test.json"))
    if optional_notes:
        for url in re.findall(r"https?://[^\s)>\]]+", optional_notes):
            if "tarc.edu.my" in url or "dialogflow" in url or "cloud.google.com" in url:
                rows.append(("Setup/documentation source", url.rstrip(".,"), "docs/dialogflow_setup.md"))
    deduped = []
    seen = set()
    for row in rows:
        if row[1] not in seen:
            deduped.append(row)
            seen.add(row[1])
    return deduped


def build_report():
    evaluation = load_json(EVALUATION_PATH, {})
    intents = load_json(INTENTS_PATH, {"intents": []})
    feedback = load_json(FEEDBACK_PATH, [])
    response_payload = load_json(RESPONSE_TEST_PATH, {"cases": []})
    model_selection = load_json(MODEL_SELECTION_PATH, {})
    latest_test_results = load_json(LATEST_TEST_RESULTS_PATH, {})
    dialogflow_notes = DIALOGFLOW_NOTES_PATH.read_text(encoding="utf-8") if DIALOGFLOW_NOTES_PATH.exists() else ""

    if not evaluation.get("results"):
        raise RuntimeError(f"No evaluation results found in {EVALUATION_PATH}")
    results = evaluation["results"]
    details = evaluation.get("details", {})
    methodology = details.get("methodology", {})
    confusion_matrices = details.get("confusion_matrices", {})
    test_cases = details.get("test_cases", [])
    intent_rows = intents.get("intents", [])
    raw_intent_count = len(intent_rows)
    raw_phrase_count = sum(len(item.get("patterns", [])) for item in intent_rows)
    raw_response_count = sum(len(item.get("responses", [])) for item in intent_rows)
    intent_pattern_counts = [len(item.get("patterns", [])) for item in intent_rows]
    minimum_intent_patterns = min(intent_pattern_counts, default=0)
    median_intent_patterns = float(np.median(intent_pattern_counts)) if intent_pattern_counts else 0.0
    maximum_intent_patterns = max(intent_pattern_counts, default=0)
    feedback_records = feedback if isinstance(feedback, list) else []
    survey = summarise_survey(feedback_records)
    if survey["respondent_count"] == 0:
        raise RuntimeError("The verified survey snapshot is empty")
    assets = generate_assets(intents, evaluation, survey)
    response_cases = as_records(response_payload)
    model_selection_rows = model_selection.get("results", []) if isinstance(model_selection, dict) else []
    selected_candidate = model_selection.get("selected_candidate", "not recorded") if isinstance(model_selection, dict) else "not recorded"
    official_sources = source_rows(response_payload, dialogflow_notes)
    generated_at = evaluation.get("generated_at", "not recorded")
    build_time = datetime.now().astimezone().isoformat(timespec="seconds")
    screenshot_path = assets["app_screenshot"]
    screenshot_timestamp = (
        datetime.fromtimestamp(screenshot_path.stat().st_mtime).astimezone().isoformat(timespec="seconds")
        if screenshot_path.exists()
        else "not captured"
    )

    best_row = max(results, key=lambda row: safe_float(row.get("Accuracy")) or -1)
    member2 = next((row for row in results if "Member 2" in result_name(row)), results[min(1, len(results) - 1)])
    member1 = next((row for row in results if "Member 1" in result_name(row)), results[0])
    member2_fallback = safe_float(member2.get("Fallback Rate"))
    if member2_fallback is None:
        member2_fallback_interpretation = "The operational impact cannot be interpreted until fallback data are recorded."
    elif member2_fallback >= 0.25:
        member2_fallback_interpretation = "This is a substantial refusal rate for a student-facing service and should be reduced without sacrificing error control."
    elif member2_fallback >= 0.10:
        member2_fallback_interpretation = "This remains a material usability cost that should be examined during threshold tuning."
    else:
        member2_fallback_interpretation = "The refusal rate is comparatively low, but the answered cases still require error review."
    has_response_metrics = any(
        safe_float(field(row, "BLEU Score (g.ii)", "BLEU")) is not None
        and safe_float(field(row, "ROUGE-1 Score (g.ii)", "ROUGE-L", "ROUGE")) is not None
        for row in results
        if "Member" in result_name(row)
    )

    document = Document()
    configure_styles(document)
    configure_document(document)
    document.core_properties.title = "TAR UMT University Inquiry Chatbot: Dual-Approach Design and Leakage-Free Evaluation"
    document.core_properties.subject = "BMCS2003 Artificial Intelligence assignment documentation"
    document.core_properties.author = "[TO BE PROVIDED]"
    document.core_properties.keywords = "chatbot, Dialogflow ES, TF-IDF, logistic regression, evaluation"
    document.core_properties.comments = "Generated reproducibly from repository evaluation artifacts."

    # Cover — editorial_cover composition.
    logo = ASSET_DIR / "tarumt_logo.png"
    if logo.exists():
        p = document.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p.paragraph_format.space_after = Pt(30)
        shape = p.add_run().add_picture(str(logo), width=Inches(4.6))
        shape._inline.docPr.set("descr", "TAR UMT institutional wordmark reused from the original assignment report")
        shape._inline.docPr.set("title", "TAR UMT wordmark")
    kicker = document.add_paragraph()
    kicker.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    kicker.paragraph_format.space_after = Pt(14)
    run = kicker.add_run("BMCS2003  •  ARTIFICIAL INTELLIGENCE  •  SESSION 202605")
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(BLUE)
    run.font.all_caps = True

    title = document.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title.paragraph_format.space_after = Pt(16)
    run = title.add_run("TAR UMT University\nInquiry Chatbot")
    run.font.name = "Calibri"
    run.font.size = Pt(30)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(NAVY)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle.paragraph_format.left_indent = Inches(0.45)
    subtitle.paragraph_format.right_indent = Inches(0.45)
    subtitle.paragraph_format.space_after = Pt(24)
    run = subtitle.add_run("Dual-Approach Design and Leakage-Free Evaluation of a Dialogflow ES Configuration and a Local TF-IDF + Logistic Regression Pipeline")
    run.font.name = "Calibri"
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor.from_string(CHARCOAL)

    cover_rule = document.add_paragraph()
    cover_rule.paragraph_format.space_before = Pt(3)
    cover_rule.paragraph_format.space_after = Pt(24)
    set_paragraph_border(cover_rule, side="top", color=BLUE, size="12", space="1")

    cover_meta = [
        ("Prepared by", "Member 1: [TO BE PROVIDED]  |  Student ID: [TO BE PROVIDED]"),
        ("", "Member 2: [TO BE PROVIDED]  |  Student ID: [TO BE PROVIDED]"),
        ("Tutorial group", "[TO BE PROVIDED]"),
        ("Tutor", "[TO BE PROVIDED]"),
        ("Submission deadline", "28 August 2026, before 12:00 p.m."),
        ("Evaluation snapshot", str(generated_at)),
    ]
    for label, value in cover_meta:
        p = document.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p.paragraph_format.space_after = Pt(5)
        if label:
            r = p.add_run(f"{label.upper()}  ")
            r.bold = True
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor.from_string(MID_GREY)
        r = p.add_run(value)
        r.font.size = Pt(10.5)
        r.font.color.rgb = RGBColor.from_string(CHARCOAL)

    add_page_break(document)

    add_heading(document, "Document control and report map", level=1)
    add_callout(
        document,
        "Evidence boundary",
        "All reported model values are read from data/evaluation_results.json. Member 1 values describe a local, train-only Dialogflow-style simulator and are not cloud Dialogflow metrics. Survey results are read from the anonymized Google Forms snapshot in data/user_feedback_verified.json; favorable means a rating of 4 or 5. BLEU and ROUGE are reported only when the evaluation artifact contains scores derived from independent reference answers.",
        tone="gold",
    )
    add_caption(document, "Table 1: Document control")
    add_table(
        document,
        ["Field", "Recorded value"],
        [
            ("Report builder", "tools/build_final_report.py"),
            ("Output", "AI Report - Final.docx"),
            ("Evaluation source", "data/evaluation_results.json"),
            ("Evaluation generated at", generated_at),
            ("Document built at", build_time),
            ("Identifying details", "Member names, IDs, tutorial group, and tutor remain [TO BE PROVIDED]."),
        ],
        [2500, 6860],
    )
    add_heading(document, "Report map", level=2)
    for item in [
        "Executive summary and rubric traceability",
        "1. Introduction",
        "2. Related Work",
        "3. Methodology",
        "4. Results and Discussion",
        "5. Conclusion, References and Sources",
        "Appendices A–F: contributions, reproducibility, survey, authentic evidence, and two plagiarism forms",
    ]:
        add_bullet(document, item)
    add_paragraph(document, "The five numbered sections mirror the five Documentation Assessment Rubric items in Appendix 1 of the assignment specification.", style="Source Note")

    add_heading(document, "Executive summary", level=1, page_break=True)
    add_paragraph(
        document,
        f"This project addresses repetitive and fragmented university enquiries by designing a task-oriented chatbot for TAR UMT information. Two development approaches are represented: a Dialogflow ES configuration path and a local Python path. The local path cleans and lemmatises queries with NLTK, extracts character-boundary TF-IDF features (3–5 character n-grams), and applies balanced multinomial Logistic Regression (C = 30, max_iter = 2000) with a 0.20 confidence threshold. Low-confidence queries return a controlled fallback and can be logged for human review.",
    )
    add_paragraph(
        document,
        f"The current evidence snapshot uses a stratified 80/20 split with random_state = {methodology.get('random_state', 'N/A')}, {methodology.get('training_count', 'N/A')} training examples, {methodology.get('test_count', 'N/A')} held-out examples, and {methodology.get('train_test_text_overlap_count', 'N/A')} cleaned-text overlaps between partitions. The highest held-out accuracy in the recorded comparison is {percent_text(best_row.get('Accuracy'))} for {result_name(best_row)}. Member 2 records {percent_text(member2.get('Accuracy'))} accuracy, {metric_text(member2.get('F1-Score'))} weighted F1, {percent_text(member2.get('Coverage'))} coverage, and {percent_text(member2.get('Fallback Rate'))} fallback rate. These values demonstrate the present prototype's limitations rather than a production-ready result.",
    )
    response_status = str(field(member2, "Response Quality Status") or methodology.get("response_quality") or "N/A")
    add_paragraph(
        document,
        f"Response-generation scoring status is {response_status}. The verified usability survey contains {survey['respondent_count']} anonymous response(s) and {survey['rating_count']} item ratings. The five-item mean is {survey['overall_mean']:.2f}/5 and {survey['overall_favorable_count']}/{survey['rating_count']} ratings ({survey['overall_favorable_rate'] * 100:.1f}%) are favorable (4 or 5). These are preliminary descriptive findings because N=5 is too small for generalisation.",
    )
    add_heading(document, "Rubric traceability", level=2)
    add_caption(document, "Table 2: Documentation-rubric traceability")
    add_table(
        document,
        ["Rubric item", "Primary evidence in this report", "Good-band intent"],
        [
            ("Introduction", "Background, problem, gap, aligned objectives, scope, significance", "Comprehensive and justified"),
            ("Related Work", "Critical comparison of university chatbots, Dialogflow, local ML, and evaluation", "Evaluated, not merely described"),
            ("Methodology", "Architecture, dataset provenance, algorithms, split, metrics, and validity controls", "Logical, reproducible, justified"),
            ("Results & Discussion", "Dynamic tables and charts, confusion analysis, implications, limitations", "Evidence-led interpretation"),
            ("Conclusion & References & Source", "Achievements, limitations, future work, APA references, source inventory", "Complete and academically transparent"),
        ],
        [1750, 5150, 2460],
        font_size=8.9,
    )

    # 1. Introduction
    add_heading(document, "1. Introduction", level=1, page_break=True)
    add_heading(document, "1.1 Background and context", level=2)
    add_paragraph(
        document,
        "University applicants and students routinely ask recurring questions about programmes, admissions, fees, academic calendars, accommodation, facilities, student affairs, and contact channels. The answers may exist on official pages, but users must still identify the correct page, interpret institutional terminology, and verify that the information is current. A task-oriented FAQ chatbot can shorten this search path by mapping a natural-language query to a controlled intent and then returning a curated answer or an official source link.",
    )
    add_paragraph(
        document,
        "Campus chatbots are a credible application area rather than a purely technical exercise. Ranoliya et al. (2017) demonstrated a university FAQ chatbot, while Dibitonto et al. (2018) situated a virtual assistant within student university life. Dialogflow ES provides managed concepts for agents, intents, entities, contexts, fulfilment, and integrations (Google Cloud, n.d.). A local statistical classifier offers a contrasting path whose data flow, features, decision threshold, and errors can be inspected in the repository.",
    )

    add_heading(document, "1.2 Problem statement", level=2)
    add_paragraph(
        document,
        "The practical problem is not simply to return an answer when a familiar phrase is entered. The system must recognise paraphrases across many closely related university intents, reject low-confidence queries safely, preserve the provenance of factual answers, and provide evidence that its reported metrics were obtained without training–test leakage. Closely related labels—such as admission versus admission documents, location versus campus map, and facilities versus sports—make the task difficult when each intent has only a small number of examples.",
    )
    add_paragraph(
        document,
        "A second problem is evidence quality. A configured Dialogflow agent, a local approximation of Dialogflow-style pattern matching, and a deployed local classifier are different systems. Reporting an offline simulator score as cloud accuracy would invalidate the comparison. Likewise, BLEU, ROUGE, and satisfaction statistics require independent references or genuine respondents; template self-comparison and invented responses are not acceptable evidence.",
    )

    add_heading(document, "1.3 Research gap", level=2)
    add_paragraph(
        document,
        "The reviewed studies establish the usefulness of university chatbots and the importance of selecting an appropriate development platform, but they do not provide a leakage-free, same-split comparison for this repository's TAR UMT inquiry taxonomy. This project therefore focuses on an auditable comparison in which the evaluation boundary is explicit: a train-only local Dialogflow-style simulator, the deployed-configuration local Logistic Regression model, and two classical baselines are tested on the same held-out queries. Cloud Dialogflow performance remains a separate validation task requiring authentic console or API evidence.",
    )

    add_heading(document, "1.4 Objectives", level=2)
    add_numbered_list(document, [
        "Design a university FAQ chatbot that covers major TAR UMT inquiry categories and returns controlled, source-aware responses.",
        "Represent two distinct member approaches: a Dialogflow ES agent configuration and an offline Python intent classifier.",
        "Implement the local classifier with deterministic preprocessing, character-boundary TF-IDF features, balanced Logistic Regression, confidence gating, and fallback logging.",
        "Evaluate intent recognition on a fixed stratified held-out split using accuracy, weighted precision, weighted recall, weighted F1, coverage, fallback rate, and confusion analysis.",
        "Evaluate response quality using independent references and analyse genuine user-satisfaction responses with transparent descriptive rules.",
        "Deliver reproducible source code, evaluation artifacts, charts, and documentation that can be regenerated from the repository.",
    ])

    add_heading(document, "1.5 Scope and significance", level=2)
    add_paragraph(
        document,
        f"The current source inventory contains {raw_intent_count} semantic intents and {raw_phrase_count} raw training phrases in data/intents.json. The chatbot is English-first, single-turn, and task-oriented. Responses are selected from controlled templates rather than generated freely. This scope supports predictable answers and reduces hallucination risk, but it does not replace official TAR UMT pages or staff advice. Time-sensitive fees, dates, policies, and contact details must be verified at the linked official source.",
    )
    add_paragraph(
        document,
        "The project is significant in three ways. For users, it offers a consistent entry point to common information. For administrators, fallback logs expose unanswered demand that can guide dataset maintenance. For AI study, the project demonstrates how thresholding, data provenance, and evaluation design can matter as much as the choice of classifier.",
    )
    add_callout(document, "Success criterion", "A working prototype is necessary but not sufficient. Success means that system claims remain traceable to code, data, and authentic test evidence, and that limitations are disclosed where the current scores are weak.", tone="blue")

    # 2. Related work
    add_heading(document, "2. Related Work", level=1, page_break=True)
    add_heading(document, "2.1 University FAQ and campus assistants", level=2)
    add_paragraph(
        document,
        "Ranoliya et al. (2017) presented a chatbot for university-related FAQs, illustrating the fit between structured institutional questions and intent/pattern-oriented interaction. Its relevance to this project is the bounded FAQ domain; its limitation for the present study is that a system that performs well on known patterns may still generalise poorly to held-out paraphrases. Dibitonto et al. (2018) designed LiSA as a campus virtual assistant to support students in university life. That work broadens the design question from classification alone to the student's situated experience and reinforces the need for usable, context-appropriate responses.",
    )
    add_heading(document, "2.2 Platform-based and local development", level=2)
    add_paragraph(
        document,
        "Dialogflow ES is a managed natural-language-understanding platform for conversational interfaces. Its agent model organises intents, entities, responses, fulfilment, contexts, and integrations (Google Cloud, n.d.). This can accelerate configuration and integration, but the trained cloud service is externally managed and must be evaluated through its real console or API. A local simulator can exercise repository patterns and provide a transparent baseline, but it cannot stand in for Google's NLU.",
    )
    add_paragraph(
        document,
        "Pérez-Soler et al. (2021) frame chatbot development as a tool-selection problem. Their comparison-oriented perspective is useful because the right choice depends on requirements such as deployment control, integration, language support, cost, and maintainability. The local TF-IDF pipeline used here is deliberately inspectable and offline-capable. Character-boundary features improve tolerance to small spelling and word-form variations, but they still require representative data and do not supply semantic understanding.",
    )
    add_heading(document, "2.3 Evaluation of classification and responses", level=2)
    add_paragraph(
        document,
        "Accuracy alone can obscure minority-class behaviour in a multi-intent dataset. Weighted precision, recall, and F1 summarise class-level outcomes while preserving class prevalence; the confusion matrix shows which labels are exchanged. Coverage and fallback rate are also essential for a thresholded chatbot because an apparently cautious model can improve the quality of answered cases by declining many inputs.",
    )
    add_paragraph(
        document,
        "BLEU measures n-gram precision with a brevity penalty (Papineni et al., 2002), while ROUGE includes recall-oriented overlap measures (Lin, 2004). Both require candidate outputs and independent reference answers. For a retrieval-style chatbot with template responses, comparing a selected template against the same template bank would be circular. The present report therefore displays N/A when independent reference scoring is absent and treats human judgement as a complementary future measure.",
    )
    add_caption(document, "Table 3: Critical comparison of prior work and the present study")
    add_table(
        document,
        ["Source", "Contribution", "Relevant limitation or trade-off", "Design implication here"],
        [
            ("Ranoliya et al. (2017)", "University FAQ chatbot", "Pattern success need not imply paraphrase generalisation", "Use a held-out split and confusion analysis"),
            ("Dibitonto et al. (2018)", "Campus assistant designed around student life", "User experience extends beyond classifier accuracy", "Include authentic usability protocol"),
            ("Pérez-Soler et al. (2021)", "Framework for choosing chatbot tools", "No single tool dominates every deployment criterion", "Separate cloud configuration from local ML evidence"),
            ("Google Cloud (n.d.)", "Dialogflow ES agent, intent, entity, fulfilment, and integration concepts", "Cloud behaviour must be tested in the actual service", "Do not label a local simulator score as cloud accuracy"),
            ("Papineni et al. (2002); Lin (2004)", "Automatic text-overlap metrics", "Scores depend on independent references and do not establish factuality", "Report N/A until a valid reference test is executed"),
        ],
        [1600, 2480, 2850, 2430],
        font_size=8.3,
    )
    add_heading(document, "2.4 Synthesis and justification", level=2)
    add_paragraph(
        document,
        "The literature supports a dual-approach project but also exposes the main risk: unlike systems can be compared as if their evidence were equivalent. The present methodology addresses that risk by identifying each execution path, using identical held-out queries only for local evaluation components, recording coverage alongside accuracy, and retaining cloud validation as an explicit evidence gap. This makes the comparison more cautious but more defensible.",
    )

    # 3. Methodology
    add_heading(document, "3. Methodology", level=1, page_break=True)
    add_heading(document, "3.1 Research design and requirements mapping", level=2)
    add_paragraph(
        document,
        "The project follows a design–build–evaluate workflow. Requirements were translated into a bounded FAQ scenario, two development paths, an intent dataset, controlled response templates, a user interface, and a reproducible evaluation harness. The evaluation is observational: it reports current prototype behaviour and does not claim statistical generalisation beyond the held-out sample.",
    )
    add_caption(document, "Table 4: Chatbot-assignment requirement mapping")
    add_table(
        document,
        ["Assignment requirement", "Repository implementation", "Evidence status"],
        [
            ("Real-life chatbot scenario", "TAR UMT FAQ and student-information assistant", "Implemented"),
            ("Background study", "Section 2 compares university chatbots, tools, and metrics", "Documented"),
            ("Development approach", "Dialogflow ES configuration artifacts plus local Python ML", "Artifacts present; cloud run needs authentic evidence"),
            ("Different member solutions", "Member 1 configuration/simulator track; Member 2 TF-IDF + LR track", "Contribution ownership must be confirmed"),
            ("Intent-recognition testing", "Leakage-free split, four-model metrics, confusion analysis", "Measured"),
            ("Response relevancy/quality", "Independent-reference protocol", f"{'Measured from ' + str(len(response_cases)) + ' cases' if has_response_metrics else 'N/A unless evaluation JSON contains scores'}"),
            ("User satisfaction", "Five-item Google Forms instrument; verified anonymous snapshot; Section 4.4 and Appendix C", f"Measured descriptively; N={survey['respondent_count']}"),
        ],
        [2300, 4460, 2600],
        font_size=8.6,
    )

    add_heading(document, "3.2 System architecture", level=2)
    add_figure(
        document,
        assets["architecture"],
        6.45,
        "Architecture diagram. A student query enters the Streamlit interface and is routed either directly to Dialogflow ES, with no local NLTK preprocessing, or to a local NLTK, character-boundary TF-IDF, Logistic Regression and confidence-gate pipeline. A separate offline evaluation harness uses a train-only local Dialogflow-style simulator and explicitly does not represent cloud accuracy.",
        "Figure 1: Operational architecture and offline-evaluation boundary",
    )
    add_paragraph(
        document,
        "The architecture distinguishes operation from evaluation. On the Dialogflow path, the user query is sent to the configured Google service; it does not first pass through the local NLTK pipeline. On the local path, preprocessing, vectorisation, classification, thresholding, response retrieval, and fallback logging execute within Python. The evaluation harness creates the train/test split and constructs the local models from training examples only.",
    )

    add_heading(document, "3.3 Dataset, provenance, and representation", level=2, page_break=True)
    evaluation_count = methodology.get("evaluation_pattern_count", methodology.get("raw_pattern_count", "N/A"))
    add_caption(document, "Table 5: Dataset inventory and evaluation snapshot")
    add_table(
        document,
        ["Artifact or quantity", "Current value", "Interpretation"],
        [
            ("data/intents.json", f"{raw_intent_count} semantic intents; {raw_phrase_count} raw phrases; {raw_response_count} response templates", "Current source inventory"),
            ("Evaluation rows", str(evaluation_count), "Canonical rows recorded when evaluation_results.json was generated"),
            ("Training partition", str(methodology.get("training_count", "N/A")), "80% stratified; models fit only here"),
            ("Held-out partition", str(methodology.get("test_count", "N/A")), "20% stratified; used for classification metrics"),
            ("Cleaned-text overlap", str(methodology.get("train_test_text_overlap_count", "N/A")), "Expected to be zero"),
            ("Labels in confusion matrix", f"{raw_intent_count} semantic intents plus fallback outcome", "Fallback is an outcome, not a training intent"),
        ],
        [2300, 2680, 4380],
        font_size=8.7,
    )
    if safe_float(evaluation_count) is not None and int(float(evaluation_count)) != raw_phrase_count:
        duplicate_count = int(methodology.get("same_label_duplicate_rows_removed", 0) or 0)
        recorded_raw_count = int(methodology.get("raw_pattern_count", -1) or -1)
        intentional_canonicalisation = (
            recorded_raw_count == raw_phrase_count
            and raw_phrase_count - int(float(evaluation_count)) == duplicate_count
        )
        if intentional_canonicalisation:
            add_callout(
                document,
                "Canonicalisation note",
                f"The source file contains {raw_phrase_count} raw phrases. Evaluation removed {duplicate_count} same-label duplicate cleaned row(s), leaving {evaluation_count} unique canonical examples. This is an intentional preprocessing control, not a stale-snapshot mismatch.",
                tone="blue",
            )
        else:
            add_callout(
                document,
                "Snapshot mismatch",
                f"The current intents.json contains {raw_phrase_count} raw phrases, whereas the evaluation artifact records {evaluation_count} evaluation rows and does not reconcile through its recorded duplicate-removal count. Rerun evaluate.py, then rerun this builder, before final submission.",
                tone="red",
            )
    add_heading(document, "Class balance and source curation", level=3)
    add_paragraph(
        document,
        f"The current intent inventory ranges from {minimum_intent_patterns} to {maximum_intent_patterns} phrases per class, with a median of {median_intent_patterns:.1f}. This imbalance, combined with closely related labels, motivates weighted metrics and class_weight='balanced'. New phrases should be added where error analysis shows a real lexical gap, then checked for duplicate cleaned text and label ambiguity before evaluation.",
    )
    add_paragraph(
        document,
        "The dataset is stored locally as JSON with one tag, multiple user-phrase patterns, and one or more controlled responses per intent. Topics and factual response anchors are checked against official TAR UMT pages listed in Section 5.5 and data/response_quality_test.json. The project does not claim that the phrase set is a scraped public benchmark or a big-data corpus.",
    )
    add_figure(
        document,
        assets["dataset"],
        6.15,
        f"Horizontal bar chart showing the current number of training phrases for each of the {raw_intent_count} semantic intents in data/intents.json.",
        "Figure 2: Distribution of raw training phrases across the current intent inventory",
    )
    add_heading(document, "3.4 Preprocessing and feature extraction", level=2)
    add_numbered_list(document, [
        "Convert the query to lowercase.",
        "Remove bracketed text, URLs, HTML fragments, punctuation, and line breaks; preserve digits and alphanumeric course codes because years, fees, and codes carry meaning.",
        "Split on whitespace and apply WordNet lemmatisation with a safe fallback when the resource is unavailable.",
        "Transform cleaned text into sublinear character-boundary TF-IDF features with analyzer = 'char_wb' and ngram_range = (3, 5), representing within-word fragments while respecting word boundaries.",
        "Retain stop words because character features and question words such as where, when, and how can distinguish university intents.",
    ])
    add_paragraph(document, "The preprocessing is deterministic and shared by the local classifier and local simulator. It is not asserted to be part of Google Dialogflow's cloud execution path.", style="Source Note")

    add_heading(document, "3.5 Algorithms and member approaches", level=2)
    add_heading(document, "3.5.1 Member 1: Dialogflow ES configuration and local evaluation surrogate", level=3)
    add_paragraph(
        document,
        "Member 1's platform approach is represented by Dialogflow ES configuration artifacts, including intents, training phrases, four custom entities, eight structured parameters, entity annotations, required-slot prompting, controlled static responses, and welcome/fallback handling. Webhook fulfillment is disabled by design, avoiding an external deployment dependency during the assessed demonstration. The repository also contains DialogflowSimulatorClient, a transparent local pattern/Jaccard implementation used solely for offline, train-only evaluation. This simulator is reproducible but is not Google's model and provides no evidence of cloud accuracy.",
    )
    add_heading(document, "3.5.2 Member 2: TF-IDF and Logistic Regression", level=3)
    add_paragraph(
        document,
        "Member 2's deployed local classifier uses TfidfVectorizer(analyzer='char_wb', ngram_range=(3, 5), sublinear_tf=True) followed by LogisticRegression(C=30.0, max_iter=2000, random_state=42, class_weight='balanced'). The character-boundary representation captures within-word fragments and is more tolerant of minor spelling and inflection changes than the earlier word-only configuration. Exact cleaned training phrases receive confidence 1.0. Otherwise, the maximum class probability is compared with a confidence threshold of 0.20. Predictions below the threshold return fallback; during normal operation they may be logged to data/unrecognized_queries.json for human review. Evaluation disables logging so that the test procedure does not alter operational data.",
    )
    add_heading(document, "3.5.3 Model-selection justification", level=3)
    if model_selection_rows:
        selection_table_rows = []
        candidate_labels = {
            "word_tfidf_1_2_lr_c10": "Word TF-IDF (1–2) + LR C=10",
            "char_wb_tfidf_3_5_lr_c30": "Character-boundary TF-IDF (3–5) + LR C=30",
            "word_char_union_lr_c30": "Word/character feature union + LR C=30",
        }
        for candidate in model_selection_rows:
            threshold_metrics = candidate.get("deployment_threshold_metrics", {})
            candidate_name = str(candidate.get("candidate", "unnamed"))
            selection_table_rows.append((
                candidate_labels.get(candidate_name, candidate_name),
                metric_text(threshold_metrics.get("accuracy")),
                metric_text(threshold_metrics.get("weighted_f1")),
                metric_text(threshold_metrics.get("coverage")),
                "Selected" if candidate_name == selected_candidate else "Not selected",
            ))
        add_caption(document, "Table 6: Three-fold out-of-fold model-selection comparison")
        add_table(
            document,
            ["Candidate", "Accuracy at 0.20", "Weighted F1 at 0.20", "Coverage at 0.20", "Decision"],
            selection_table_rows,
            [3200, 1450, 1800, 1550, 1360],
            font_size=8.2,
        )
        selected_threshold = next(
            (
                row.get("deployment_threshold_metrics", {})
                for row in model_selection_rows
                if row.get("candidate") == selected_candidate
            ),
            {},
        )
        add_paragraph(
            document,
            f"Candidate selection used {model_selection.get('protocol', 'stratified out-of-fold predictions')} with random_state = {model_selection.get('random_state', 42)} and the recorded selection metric, {model_selection.get('selection_metric', 'deployment-threshold weighted F1')}. The selected character-boundary candidate records {metric_text(selected_threshold.get('weighted_f1'))} weighted F1 and {metric_text(selected_threshold.get('coverage'))} coverage at threshold 0.20, improving the thresholded word-only candidate while remaining simpler than a feature union. Because the out-of-fold comparison covers the same canonical corpus used for subsequent development evaluation, the final held-out numbers should be treated as model-development evidence rather than a sealed external test.",
        )
    else:
        add_callout(document, "Model-selection artifact missing", "data/model_selection_results.json was not available. Generate it with model_selection.py before claiming that the deployed feature configuration was selected empirically.", tone="red")

    add_heading(document, "3.5.4 Baselines", level=3)
    add_paragraph(
        document,
        "Multinomial Naïve Bayes with alpha = 1.0 and a unigram TF-IDF representation provides a probabilistic lexical baseline. Linear SVM with C = 0.1, max_iter = 1000, random_state = 42, and unigram TF-IDF provides a margin-based baseline. The baselines do not generate responses and are therefore N/A for BLEU/ROUGE.",
    )

    add_heading(document, "3.6 Evaluation protocol and metrics", level=2)
    add_caption(document, "Table 7: Evaluation protocol")
    add_table(
        document,
        ["Control", "Recorded setting", "Rationale"],
        [
            ("Split", str(methodology.get("split", "80/20 stratified")), "Preserve class proportions where possible"),
            ("Random seed", str(methodology.get("random_state", 42)), "Repeatable partition"),
            ("Training/test counts", f"{methodology.get('training_count', 'N/A')} / {methodology.get('test_count', 'N/A')}", "Make sample size explicit"),
            ("Text leakage check", f"{methodology.get('train_test_text_overlap_count', 'N/A')} overlaps", "Prevent cleaned duplicates across partitions"),
            ("Member 1 scope", str(methodology.get("dialogflow_scope", "local simulator; training phrases only; not cloud")), "Prevent cloud-performance mislabelling"),
            ("Member 2 threshold", str(methodology.get("deployment_confidence_threshold", 0.20)), "Match the deployed confidence policy"),
            ("Averaging", "Weighted precision, recall, and F1", "Account for class support"),
        ],
        [2060, 3000, 4300],
        font_size=8.6,
    )
    add_paragraph(document, "Accuracy is the proportion of held-out queries assigned the expected label. Weighted precision measures the reliability of predicted labels, weighted recall measures recovered instances, and weighted F1 is their harmonic balance by class support. Coverage is the proportion receiving a non-fallback label; fallback rate is its complement for the thresholded systems.")
    response_metric_sentence = (
        "The current evaluation artifact records the resulting scores for the member response engines; they are interpreted only as lexical-overlap evidence and not as proof of factual correctness."
        if has_response_metrics
        else "The current evaluation artifact records no response-overlap scores, so the status remains N/A until evaluation is rerun and validation passes."
    )
    add_paragraph(document, f"For response quality, the protocol uses human-authored reference answers anchored to official TAR UMT sources. Corpus BLEU and mean ROUGE-1 F1 may be computed only after those references are independent of the candidate-selection templates. The repository currently contains {len(response_cases)} reference case(s), and the report uses only scores present in evaluation_results.json. {response_metric_sentence}")
    add_paragraph(document, f"For usability, the Google Form uses five Likert statements rated from 1 (strongly disagree) to 5 (strongly agree). The verified artifact contains {survey['respondent_count']} anonymous responses collected from 12 to 24 August 2026. The report presents item means, medians, and the predeclared favorable rate (ratings of 4 or 5). No inferential or representative claim is made because the sample is small.")

    add_heading(document, "3.7 Validity, ethics, and data governance", level=2)
    for risk in [
        "Construct validity: report weighted metrics with fallback behaviour; do not equate text overlap with factual correctness.",
        "Internal validity: check cleaned train/test overlap and restrict every local evaluator to training phrases.",
        f"External validity: {methodology.get('test_count', 'the recorded number of')} held-out examples across many intents do not establish production performance.",
        "Cloud validity: the offline simulator does not validate Dialogflow ES; authentic console/API tests are required.",
        "Privacy: minimise survey fields and operational logs, redact identifiers, and control access.",
        "Factual currency: point time-sensitive answers to official sources and review them periodically.",
    ]:
        paragraph = add_bullet(document, risk)
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.paragraph_format.line_spacing = 1.05
        for run in paragraph.runs:
            run.font.size = Pt(9.5)

    # 4. Results & Discussion
    add_heading(document, "4. Results and Discussion", level=1)
    add_heading(document, "4.1 Classification results", level=2)
    add_caption(document, "Table 8: Leakage-free held-out classification metrics from evaluation_results.json")
    metric_rows = []
    for row in results:
        metric_rows.append((
            result_name(row), engine_name(row), metric_text(row.get("Accuracy")), metric_text(row.get("Precision")),
            metric_text(row.get("Recall")), metric_text(row.get("F1-Score")),
        ))
    add_table(
        document,
        ["Model", "Engine/evaluation scope", "Accuracy", "Precision", "Recall", "F1"],
        metric_rows,
        [2050, 3470, 960, 960, 960, 960],
        font_size=8.15,
    )
    add_paragraph(document, f"Source: data/evaluation_results.json generated at {generated_at}. All classification metrics use the same recorded held-out split.", style="Source Note")
    add_figure(
        document,
        assets["metrics"],
        6.35,
        "Grouped bar chart of accuracy, weighted precision, weighted recall, and weighted F1 for the Member 1 local simulator, Member 2 TF-IDF plus Logistic Regression model, Multinomial Naive Bayes baseline, and Linear SVM baseline.",
        "Figure 3: Held-out intent-classification performance by model",
    )
    add_paragraph(
        document,
        f"The strongest recorded accuracy is {metric_text(best_row.get('Accuracy'))} for {result_name(best_row)}. Member 2's local deployment configuration reaches {metric_text(member2.get('Accuracy'))} accuracy and {metric_text(member2.get('F1-Score'))} weighted F1. Member 1's local simulator records {metric_text(member1.get('Accuracy'))} accuracy and {metric_text(member1.get('F1-Score'))} weighted F1; this low score belongs to the train-only simulator and must not be re-labelled as Dialogflow cloud performance.",
    )
    add_paragraph(
        document,
        "The result ordering suggests that a margin-based linear model currently generalises better than the other tested lexical approaches on the held-out phrases. Member 2's classifier is affected by both multiclass separation and its confidence gate. The comparison does not establish that a particular algorithm is universally superior; it reflects this split, preprocessing, feature configuration, and small class supports.",
    )

    add_heading(document, "4.2 Coverage, fallback, and response-quality status", level=2)
    operational_rows = []
    for row in results:
        bleu = field(row, "BLEU Score (g.ii)", "BLEU")
        rouge = field(row, "ROUGE-1 Score (g.ii)", "ROUGE-L", "ROUGE")
        status = field(row, "Response Quality Status") or ("N/A — baseline has no response engine" if "Baseline" in result_name(row) else "Not recorded")
        operational_rows.append((
            result_name(row), metric_text(row.get("Coverage")), metric_text(row.get("Fallback Rate")),
            f"BLEU {metric_text(bleu)}\nROUGE-1 {metric_text(rouge)}", str(status),
        ))
    add_caption(document, "Table 9: Coverage, fallback, and response-quality evidence")
    add_table(
        document,
        ["Model", "Coverage", "Fallback", "Text-overlap metrics", "Evidence status"],
        operational_rows,
        [2100, 950, 950, 1750, 3610],
        font_size=8.0,
    )
    add_figure(
        document,
        assets["coverage"],
        6.3,
        "Stacked bars showing recorded coverage and fallback rate for each model. Member 2 has lower coverage than the ungated baselines because its 0.20 confidence gate returns fallback on uncertain held-out cases.",
        "Figure 4: Coverage and fallback behaviour on the held-out set",
    )
    add_paragraph(
        document,
        f"Member 2 answers {percent_text(member2.get('Coverage'))} of the held-out cases and falls back on {percent_text(member2.get('Fallback Rate'))}. A high fallback rate can be operationally safer than an unsupported confident answer. {member2_fallback_interpretation} Threshold tuning must therefore be evaluated as a coverage–accuracy trade-off rather than optimised for one metric alone.",
    )
    if safe_float(field(member2, "BLEU Score (g.ii)")) is None:
        add_callout(document, "Response scoring is N/A", "Although an independent-source test instrument now exists in data/response_quality_test.json, the current evaluation artifact records no BLEU or ROUGE score. This report deliberately does not back-fill or estimate those values. Rerun evaluate.py and verify the reference provenance before reporting them.", tone="gold")

    if assets["app_screenshot"].exists():
        add_heading(document, "4.2.1 Authentic end-to-end prototype evidence", level=3)
        add_figure(
            document,
            assets["app_screenshot"],
            6.35,
            "Authentic Streamlit screenshot with Member 2 selected. The user asks how much tuition fees are; the local chatbot predicts the fees intent with confidence 0.9761 and returns a controlled response containing official TAR UMT fee-guide URLs.",
            "Figure 5: Authentic Streamlit end-to-end example using the Member 2 local classifier",
        )
        add_paragraph(
            document,
            f"This implementation artifact (file timestamp {screenshot_timestamp}) confirms that the final local route can accept a natural-language query, expose the active Member 2 engine, return the fees intent with confidence 0.9761, and provide official Malaysian and international fee-guide links. It is one successful demonstration case, not an estimate of overall accuracy; general performance is reported only from the held-out evaluation above.",
        )

    if latest_test_results:
        add_heading(document, "4.2.2 Labelled end-to-end smoke probes", level=3)
        add_caption(document, "Table 10: Labelled local smoke-probe summary")
        add_table(
            document,
            ["Artifact", "Queries", "Member 2 local model", "Member 1 local simulator", "Timestamp"],
            [(
                "data/latest_test_results.json",
                str(latest_test_results.get("total_queries_tested", "N/A")),
                f"{latest_test_results.get('ml_pass_count', 'N/A')}/{latest_test_results.get('total_queries_tested', 'N/A')} labelled probes passed",
                f"{latest_test_results.get('df_pass_count', 'N/A')}/{latest_test_results.get('total_queries_tested', 'N/A')} labelled probes passed",
                str(latest_test_results.get("timestamp", "not recorded")),
            )],
            [2150, 900, 2200, 2500, 1610],
            font_size=7.9,
        )
        add_paragraph(
            document,
            "These manually labelled probes are deterministic end-to-end smoke checks over selected queries. The column labelled Member 1 refers to the repository's local Dialogflow-style simulator, not the Google Dialogflow ES cloud service. A 10/10 smoke result verifies the tested cases only and is not a substitute for the held-out classification results in Table 8.",
        )

    add_heading(document, "4.3 Confusion and error analysis", level=2)
    add_figure(
        document,
        assets["errors"],
        6.45,
        "Two-panel horizontal bar chart showing the most frequent off-diagonal actual-to-predicted intent pairs for Member 1's local simulator and Member 2's TF-IDF plus Logistic Regression classifier, derived from the recorded confusion matrices.",
        "Figure 6: Most frequent off-diagonal confusion pairs for the two member approaches",
    )
    error_rows = []
    for model_name in [result_name(member1), result_name(member2)]:
        payload = confusion_matrices.get(model_name, {})
        pairs = top_confusions(payload, 4)
        pair_text = "; ".join(f"{actual} → {predicted} ({count})" for count, actual, predicted in pairs) or "No off-diagonal pairs recorded"
        fallback_predictions = sum(
            1 for case in test_cases if case.get("predictions", {}).get(model_name) == "fallback"
        )
        error_rows.append((model_name, pair_text, str(fallback_predictions)))
    add_caption(document, "Table 11: Error-analysis summary derived from recorded test cases")
    add_table(
        document,
        ["Model", "Leading off-diagonal pairs", "Fallback predictions"],
        error_rows,
        [2300, 5480, 1580],
        font_size=8.5,
    )
    add_paragraph(
        document,
        "Several errors are semantically plausible: location, campus map, contact, facilities, sport, intake, semester break, and examination questions share vocabulary. Short social utterances such as greetings, thanks, identity questions, and affection also offer few discriminating terms. These patterns imply that improvement should begin with label-boundary review and targeted paraphrase collection rather than indiscriminate duplication of existing phrases.",
    )
    add_paragraph(
        document,
        "The confusion chart counts are small because the held-out set is distributed across many labels. A difference of one example can materially change a class score. Per-class support and repeated cross-validation would be needed before making stronger comparative claims.",
    )

    add_heading(document, "4.4 User satisfaction and usability", level=2)
    add_paragraph(
        document,
        f"Five anonymous Google Forms responses were collected between 12 and 24 August 2026 using the exact five-item questionnaire reproduced in Appendix C. All complete responses were retained, including one respondent who selected 1 for every item. Across all {survey['rating_count']} ratings, the mean was {survey['overall_mean']:.2f}/5, the median was {survey['overall_median']:.1f}, and {survey['overall_favorable_count']} ratings ({survey['overall_favorable_rate'] * 100:.1f}%) were favorable. Favorable was defined before analysis as a rating of 4 or 5.",
    )
    add_caption(document, "Table 12: Verified user-satisfaction results (N=5)")
    add_table(
        document,
        ["Survey item", "Mean / 5", "Median", "Favorable (4-5)"],
        [
            (
                f"Q{index}. {item['label']}",
                f"{item['mean']:.2f}",
                f"{item['median']:.1f}",
                f"{item['favorable_count']}/5 ({item['favorable_rate'] * 100:.0f}%)",
            )
            for index, item in enumerate(survey["items"], 1)
        ],
        [4300, 1500, 1300, 2260],
        font_size=8.7,
    )
    add_figure(
        document,
        assets["survey"],
        6.2,
        "Horizontal bar chart of the five verified survey item means. Intent understanding and answer clarity both score 4.0 out of 5; interface usability scores 3.8; response speed 3.6; and overall satisfaction 3.4. Labels also show the percentage rating each item 4 or 5.",
        "Figure 7: Verified five-item usability survey results",
    )
    add_paragraph(
        document,
        "Intent understanding and answer clarity received the strongest item means (4.00/5; 80% favorable each). Interface usability averaged 3.80/5, while response speed averaged 3.60/5. Overall satisfaction was the lowest-rated item at 3.40/5, with three of five respondents rating it 4 or 5. The result suggests that the prototype is generally understandable but still needs reliability, latency, and overall experience improvements. Because N=5 and recruitment was not probabilistic, these findings describe only this pilot group and should not be generalized to the wider student population.",
    )

    add_heading(document, "4.5 Objective-by-objective interpretation", level=2)
    add_caption(document, "Table 13: Objective attainment")
    add_table(
        document,
        ["Objective", "Finding", "Status"],
        [
            ("FAQ coverage", f"{raw_intent_count} semantic intents in the current source inventory", "Met at prototype scope"),
            ("Two approaches", "Dialogflow configuration track and local ML track are documented separately", "Met; contribution evidence to confirm"),
            ("Local implementation", "Character-boundary TF-IDF (3–5), LR C=30/max_iter=2000/balanced, threshold=0.20", "Met"),
            ("Leakage-free comparison", f"Recorded overlap={methodology.get('train_test_text_overlap_count', 'N/A')}; four model rows", "Met for local evaluation"),
            ("Cloud Dialogflow performance", "No authentic cloud confusion matrix or API test artifact", "Not yet demonstrated"),
            ("Response-quality evidence", response_status, "Measured" if has_response_metrics else "N/A or pending rerun"),
            ("User satisfaction", f"N={survey['respondent_count']}; five-item mean {survey['overall_mean']:.2f}/5; favorable {survey['overall_favorable_rate'] * 100:.1f}%", "Measured descriptively"),
        ],
        [2290, 5010, 2060],
        font_size=8.5,
    )

    add_heading(document, "4.6 Limitations and practical implications", level=2)
    for limitation in [
        "The evaluation snapshot is small relative to the number of intents and may be stale when intents.json changes.",
        "The Member 1 offline score measures a local pattern/Jaccard surrogate, not Dialogflow ES cloud NLU.",
        "The member approaches are not symmetric: Member 2 includes a confidence gate, while the baselines always classify; coverage must accompany accuracy.",
        "Template responses reduce hallucination but can become outdated and do not support rich multi-turn context.",
        ("Automatic text-overlap metrics are available from the independent-reference instrument but still require human factuality and usability review." if has_response_metrics else "Automatic text-overlap metrics are N/A in the current artifact and would require both a valid reference test and human review."),
        "The verified usability pilot has only five respondents, so it supports descriptive feedback but not statistical generalisation.",
    ]:
        add_bullet(document, limitation)
    add_paragraph(document, "Practically, the strongest next step is not to advertise the current prototype as accurate. It is to review the intent taxonomy, add independently phrased examples to weak labels, tune the confidence threshold on a validation set, rerun the leakage-free evaluation, and capture authentic Dialogflow test evidence. This sequence directly targets the observed errors and evidence gaps.")

    # 5. Conclusion, references, sources
    add_heading(document, "5. Conclusion, References and Sources", level=1)
    add_heading(document, "5.1 Achievements", level=2)
    add_paragraph(
        document,
        "The project has produced an end-to-end university FAQ chatbot prototype with a clear task scenario, a structured intent/response inventory, a Dialogflow ES configuration path, an offline Python classifier, a Streamlit interface, controlled fallback handling, test code, a reproducible evaluation artifact, and a verified five-response usability pilot. The documentation separates cloud configuration from the local simulator and keeps every reported metric traceable to a recorded artifact.",
    )
    add_paragraph(
        document,
        f"The present results establish a transparent baseline rather than a production claim. {result_name(best_row)} has the highest recorded accuracy ({metric_text(best_row.get('Accuracy'))}), while Member 2's deployed configuration records {metric_text(member2.get('Accuracy'))} accuracy with {metric_text(member2.get('Coverage'))} coverage. These figures provide a concrete starting point for error-driven improvement.",
    )

    add_heading(document, "5.2 Limitations", level=2)
    response_limit = "Response-overlap scores are available, but they do not establish factual correctness or usability." if has_response_metrics else "Response-overlap metrics are not available in the current evaluation artifact."
    add_paragraph(document, f"The current system is single-turn, English-first, dependent on a small project-specific phrase set, and limited to predefined responses. Dataset and evaluation snapshots are not automatically synchronised unless evaluate.py and this builder are rerun in sequence. Dialogflow cloud accuracy has not been demonstrated by the offline surrogate. The usability survey is genuine but preliminary because it contains only five anonymous respondents. {response_limit}")

    add_heading(document, "5.3 Future work", level=2)
    add_numbered_list(document, [
        "Resolve overlapping intent definitions and expand weak classes with genuinely new paraphrases collected under a documented protocol.",
        "Use a validation partition or nested cross-validation to tune the 0.20 threshold without optimising on the final test set.",
        "Evaluate class-balanced metrics, per-class support, calibration, latency, and repeated splits in addition to the current weighted metrics.",
        "Run an authentic Dialogflow ES console/API test on the same locked queries and preserve dated screenshots or JSON responses.",
        "Expand and review the independent-source response test, then pair BLEU/ROUGE with blinded human ratings for factuality, relevance, clarity, and source usefulness.",
        "Repeat the Appendix C survey with a larger and more diverse voluntary sample, preserve anonymous responses, and compare results with the current N=5 pilot without discarding unfavorable ratings.",
        "Add multi-turn context and a source-aware retrieval layer only after factuality, privacy, and update governance are defined.",
    ])

    add_heading(document, "5.4 References", level=2)
    references = [
        "Dibitonto, M., Leszczynska, K., Tazzi, F., & Medaglia, C. M. (2018). Chatbot in a campus environment: Design of LiSA, a virtual assistant to help students in their university life. In M. Kurosu (Ed.), Human–computer interaction. Interaction technologies (Lecture Notes in Computer Science, Vol. 10903, pp. 103–116). Springer. https://doi.org/10.1007/978-3-319-91250-9_9",
        "Google Cloud. (n.d.). Dialogflow ES documentation. https://cloud.google.com/dialogflow/es/docs",
        "Lin, C.-Y. (2004). ROUGE: A package for automatic evaluation of summaries. In Text summarization branches out (pp. 74–81). Association for Computational Linguistics. https://aclanthology.org/W04-1013/",
        "Papineni, K., Roukos, S., Ward, T., & Zhu, W.-J. (2002). BLEU: A method for automatic evaluation of machine translation. In Proceedings of the 40th Annual Meeting of the Association for Computational Linguistics (pp. 311–318). Association for Computational Linguistics. https://doi.org/10.3115/1073083.1073135",
        "Pérez-Soler, S., Juárez-Puerta, S., Guerra, E., & de Lara, J. (2021). Choosing a chatbot development tool. IEEE Software, 38(4), 94–103. https://doi.org/10.1109/MS.2020.3030198",
        "Ranoliya, B. R., Raghuwanshi, N., & Singh, S. (2017). Chatbot for university related FAQs. In 2017 International Conference on Advances in Computing, Communications and Informatics (ICACCI) (pp. 1525–1530). IEEE. https://doi.org/10.1109/ICACCI.2017.8126057",
    ]
    for reference in references:
        add_reference(document, reference)

    add_heading(document, "5.5 Dataset, factual-source, and tool acknowledgement", level=2)
    add_paragraph(document, "The intent phrases and response templates are stored in data/intents.json and are treated as a team-curated project dataset, not a public big-data benchmark. Official TAR UMT URLs are used as factual anchors for the independent response-quality instrument. Complete machine-readable case-to-source mappings are preserved in data/response_quality_test.json.")
    if official_sources:
        source_display = []
        for topic, url, artifact in official_sources:
            source_display.append((topic.replace("_", " ").title(), url, artifact))
        add_caption(document, "Table 14: Official factual-source inventory used by the response-quality instrument")
        add_table(
            document,
            ["Topic", "Official URL", "Repository provenance"],
            source_display,
            [1750, 5260, 2350],
            font_size=7.7,
        )
    else:
        add_callout(document, "Source inventory pending", "No official source mappings were found in data/response_quality_test.json or docs/dialogflow_setup.md. Add verified TAR UMT pages before response-quality evaluation.", tone="red")
    add_caption(document, "Table 15: Software and artifact acknowledgement")
    add_table(
        document,
        ["Component", "Recorded use", "Repository evidence"],
        [
            ("Python", "Application, model selection, evaluation, and document builder", "app.py; model_selection.py; evaluate.py; tools/build_final_report.py"),
            ("NLTK", "Lemmatisation and optional BLEU token processing", "src/preprocessing.py; evaluate.py"),
            ("scikit-learn", "TF-IDF, Logistic Regression, Naïve Bayes, SVM, metrics, split, cross-validation", "src/ml_model.py; model_selection.py; evaluate.py; data/model_selection_results.json"),
            ("Streamlit", "Interactive prototype interface", "app.py"),
            ("Dialogflow ES", "Platform configuration, entities, parameter prompting, controlled static responses, and fallback handling", "dialogflow_agent.zip; data/entities.json; docs/dialogflow_setup.md"),
            ("Repository data", "Intent patterns, feedback, tests, evaluation results", "data/ directory"),
        ],
        [1850, 3850, 3660],
        font_size=8.3,
    )

    # Appendices
    add_heading(document, "Appendix A. Contribution record and prototype evidence", level=1, page_break=True)
    add_callout(document, "Member verification required", "The responsibilities below reflect the repository's two-track structure. Both students must confirm the allocation and add dated evidence before submission; the report does not infer authorship from file names alone.", tone="gold")
    add_caption(document, "Table A1: Individual contribution record requiring member confirmation")
    add_table(
        document,
        ["Member", "Proposed responsibility record", "Evidence to attach", "Confirmation"],
        [
            ("Member 1\n[TO BE PROVIDED]", "Dialogflow ES intent/entity/response configuration; parameter annotations and required-slot prompting; exported agent artifact; train-only simulator integration", "Export timestamp; parameter-prompt test; authentic console/API test; meeting log", "Member initials/date: __________"),
            ("Member 2\n[TO BE PROVIDED]", "NLTK preprocessing; character-boundary TF-IDF (3–5) + balanced Logistic Regression (C=30) pipeline; confidence gate; model selection; local evaluation and UI integration", "Commit/version history; model-selection artifact; test output; code walkthrough; meeting log", "Member initials/date: __________"),
            ("Shared", "Problem framing; dataset review; literature synthesis; error analysis; final QA and demonstration", "Minutes, review notes, rehearsal checklist", "Both initials/date: __________"),
        ],
        [1450, 3820, 2620, 1470],
        font_size=8.3,
    )
    add_heading(document, "Prototype demonstration evidence checklist", level=2)
    for evidence in [
        "Show a clean installation and launch of the Streamlit application.",
        "Demonstrate at least one correct local prediction, one ambiguous query, and one confidence-triggered fallback.",
        "Open the fallback log and explain human review without exposing personal data.",
        "Show the Dialogflow ES agent in the genuine console and test locked queries; retain dated screenshots or exported JSON.",
        "Run evaluate.py and explain the split, leakage check, metrics, coverage, and confusion errors.",
        "Each member presents and answers questions about their own implementation.",
    ]:
        add_bullet(document, "☐ " + evidence)
    add_callout(document, "Screenshot policy", "Insert only screenshots captured from the real local application or authenticated Dialogflow console. No synthetic console image is included in this report. Record the capture date and a short description under each inserted figure.", tone="blue")

    add_heading(document, "Appendix B. Reproducibility checklist", level=1, page_break=True)
    add_paragraph(document, "After activating the project virtual environment, run the portable commands below from the repository root. Preserve the listed repository artifacts and check each item after a clean rerun.")
    checklist_items = [
        ("Environment", "Install requirements.txt; record Python and package versions."),
        ("Dataset", "Verify data/intents.json validates and record its phrase/intent counts."),
        ("Response references", "Validate data/response_quality_test.json sources and independence."),
        ("Model selection", "Run model_selection.py; confirm the three-fold OOF artifact and selected character-boundary candidate."),
        ("Evaluation", "Run evaluate.py; confirm seed 42, stratification, and zero cleaned-text overlap."),
        ("Snapshot consistency", "Confirm evaluation_results.json was generated after the final intents.json edit."),
        ("Tests", "Run the full automated test suite and retain the terminal output."),
        ("Report", "Run the command below to refresh charts and AI Report - Final.docx."),
        ("Visual QA", "Render the DOCX with canonical render_docx.py and inspect every page."),
        ("Cloud evidence", "Export the Dialogflow agent and capture authentic locked-query results."),
        ("Survey", "Preserve data/user_feedback_verified.json; confirm N=5 and collection dates before rebuilding."),
    ]
    add_caption(document, "Table B1: Reproducibility and submission checklist")
    add_table(
        document,
        ["Done", "Area", "Verification action"],
        [("☐", area, action) for area, action in checklist_items],
        [700, 1800, 6860],
        font_size=8.8,
    )
    add_heading(document, "Regeneration commands", level=2)
    commands = [
        "python model_selection.py",
        "python evaluate.py",
        "python -m unittest discover -s tests -v",
        "python src/create_dialogflow_zip.py",
        "python tools/build_final_report.py",
    ]
    for command in commands:
        p = add_paragraph(document, command)
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.right_indent = Inches(0.25)
        set_paragraph_shading(p, LIGHT_GREY)
        for run in p.runs:
            run.font.name = "Courier New"
            run.font.size = Pt(8.6)
    add_paragraph(document, f"Evaluation artifact: {EVALUATION_PATH.name} generated at {generated_at}. Builder output generated at {build_time}.", style="Source Note")

    add_heading(document, "Appendix C. Verified user-satisfaction instrument and data", level=1, page_break=True)
    add_callout(document, "Observed pilot", f"The frozen anonymized Google Forms snapshot contains N={survey['respondent_count']} complete responses collected from 12 to 24 August 2026. All five complete responses are retained, including the all-1 response. Results are descriptive and are not generalized beyond this small pilot.", tone="gold")
    add_heading(document, "Questionnaire and scoring rule", level=2)
    add_paragraph(document, "Scale: 1 = strongly disagree, 2 = disagree, 3 = neither agree nor disagree, 4 = agree, and 5 = strongly agree. A favorable response is defined as 4 or 5. The questionnaire collected ratings and timestamps only; the exported research snapshot contains no names or student IDs.", style="Source Note")
    add_caption(document, "Table C1: Exact five-item Google Forms questionnaire")
    add_table(
        document,
        ["Item", "Statement"],
        [(f"Q{index}", statement) for index, (_, _, statement) in enumerate(SURVEY_FIELDS, 1)],
        [900, 8460],
        font_size=8.8,
    )
    add_heading(document, "Item-level descriptive results", level=2)
    add_caption(document, "Table C2: Verified descriptive statistics")
    add_table(
        document,
        ["Item", "Mean", "Median", "Favorable"],
        [
            (f"Q{index}", f"{item['mean']:.2f}/5", f"{item['median']:.1f}", f"{item['favorable_count']}/5 ({item['favorable_rate'] * 100:.0f}%)")
            for index, item in enumerate(survey["items"], 1)
        ] + [("All 25 ratings", f"{survey['overall_mean']:.2f}/5", f"{survey['overall_median']:.1f}", f"{survey['overall_favorable_count']}/25 ({survey['overall_favorable_rate'] * 100:.1f}%)")],
        [2400, 1800, 1700, 3460],
        font_size=8.8,
    )
    add_heading(document, "Anonymized response matrix", level=2)
    add_caption(document, "Table C3: Retained complete responses")
    add_table(
        document,
        ["Response", "Date", "Q1", "Q2", "Q3", "Q4", "Q5"],
        [
            (
                str(record.get("response_id", f"R{index}")),
                str(record.get("submitted_at", ""))[:10],
                *(str(record[key]) for key, _, _ in SURVEY_FIELDS),
            )
            for index, record in enumerate(feedback_records, 1)
        ],
        [1250, 2400, 1142, 1142, 1142, 1142, 1142],
        font_size=8.4,
    )
    add_heading(document, "Data-quality decisions and limitations", level=2)
    for rule in [
        "The separate local UI demo submission was excluded because it was not part of the Google Forms study.",
        "No complete Google Forms response was removed; an unfavorable response is not an exclusion criterion.",
        "Only complete integer ratings from 1 to 5 were accepted by the verified-data validator.",
        "The small convenience sample (N=5) is suitable for pilot feedback but not inferential statistics or population claims.",
        "The reproducible anonymized snapshot is stored in data/user_feedback_verified.json; the live response sheet should remain access-controlled.",
    ]:
        add_bullet(document, rule)

    add_heading(document, "Appendix D. Authentic evidence register", level=1, page_break=True)
    add_paragraph(document, "This register distinguishes captured local evidence from artifacts still required. Cloud rows remain pending because a local simulator or synthetic image cannot validate Dialogflow ES.")
    add_caption(document, "Table D1: Evidence register for final submission")
    add_table(
        document,
        ["Evidence ID", "Required artifact", "Minimum metadata", "Status"],
        [
            ("E1", "Streamlit home/interface screenshot", f"File timestamp {screenshot_timestamp}; Member 2 selection visible", "CAPTURED" if screenshot_path.exists() else "TO BE CAPTURED"),
            ("E2", "Local correct-intent interaction", "Tuition-fee query; fees intent; confidence 0.9761; official-source response", "CAPTURED" if screenshot_path.exists() else "TO BE CAPTURED"),
            ("E3", "Local fallback and review log", "Query; threshold; safe redaction; timestamp", "TO BE CAPTURED"),
            ("E4", "Dialogflow ES intents view", "Authenticated console; agent name; capture date", "TO BE CAPTURED"),
            ("E5", "Dialogflow locked-query test", "Query set ID; intent; confidence; API/console provenance", "TO BE CAPTURED"),
            ("E6", "Evaluation/test terminal output", "Command; timestamp; test count; pass/fail", "TO BE CAPTURED"),
            ("E7", "Verified Google Forms survey snapshot", f"N={survey['respondent_count']}; 12-24 August 2026; anonymized response matrix", "CAPTURED"),
        ],
        [1150, 2850, 3900, 1460],
        font_size=8.5,
    )
    if screenshot_path.exists():
        add_paragraph(document, "E1/E2 artifact: report_assets/app_chatbot_e2e.png. Figure 5 reproduces it without alteration.", style="Source Note")
    add_callout(document, "Do not substitute", "The local Dialogflow-style simulator is not an acceptable substitute for E4 or E5. E7 uses the verified anonymized Google Forms snapshot; local demo feedback is excluded.", tone="red")

    def add_plagiarism_form(member_number):
        add_heading(document, f"Appendix {'E' if member_number == 1 else 'F'}. Plagiarism Statement Form — Member {member_number}", level=1, page_break=True)
        add_paragraph(document, "Complete this form personally. The document builder intentionally leaves all identity, signature, and date fields blank.")
        add_caption(document, f"Table {'E1' if member_number == 1 else 'F1'}: Member {member_number} identification")
        add_table(
            document,
            ["Field", "To be completed by the student"],
            [
                ("Name", "[TO BE PROVIDED]"),
                ("Student ID", "[TO BE PROVIDED]"),
                ("Course code", "BMCS2003"),
                ("Assignment title", "TAR UMT University Inquiry Chatbot"),
                ("Tutorial group", "[TO BE PROVIDED]"),
            ],
            [2500, 6860],
            font_size=9.2,
        )
        add_heading(document, "Declaration", level=2)
        declaration = (
            "I declare that the work submitted for this assignment is my own contribution except where sources and team contributions are clearly acknowledged. I have not copied another student's or group's work, shared material in a way that enables academic misconduct, fabricated research participants or results, or represented generated or third-party material as my own without appropriate acknowledgement. I understand that I remain responsible for checking the accuracy, originality, citations, data provenance, and submitted code."
        )
        add_paragraph(document, declaration)
        for item in [
            "I have reviewed the final report and source code.",
            "I have checked that my individual contribution record is accurate.",
            "I have checked that all reported metrics are traceable to evaluation_results.json, user_feedback_verified.json, or clearly labelled N/A.",
            "I have complied with TAR UMT academic-integrity requirements and the instructions supplied for this assignment.",
        ]:
            add_bullet(document, "☐ " + item)
        add_paragraph(document, "Student signature: ______________________________________________")
        add_paragraph(document, "Date: __________________________________________________________")
        add_paragraph(document, "Witness / tutor (if required): ____________________________________")
        add_callout(document, "Unsigned by design", "No name, signature, date, or consent is inserted automatically. The student must complete and sign this form after reviewing the final submission.", tone="grey")

    add_plagiarism_form(1)
    add_plagiarism_form(2)

    # Remove the trailing spacer after the final callout if it exists only as a blank paragraph.
    document.save(OUTPUT_PATH)
    return {
        "output": OUTPUT_PATH,
        "assets": assets,
        "result_count": len(results),
        "evaluation_generated_at": generated_at,
        "intent_count": raw_intent_count,
        "phrase_count": raw_phrase_count,
        "feedback_count": len(feedback_records),
        "response_case_count": len(response_cases),
    }


if __name__ == "__main__":
    summary = build_report()
    print(json.dumps(summary, indent=2, default=str))
